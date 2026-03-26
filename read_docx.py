from docx import Document

doc = Document(r'c:\Users\27813\OneDrive - Konecta Pty Ltd\Software Development\Konecta Website\Konecta website additional content.docx')

for i, para in enumerate(doc.paragraphs):
    style = para.style.name if para.style else ''
    text = para.text.strip()
    if text:
        print(f'[{i}][{style}] {text}')

for table_idx, table in enumerate(doc.tables):
    print(f'\n=== TABLE {table_idx} ===')
    for row_idx, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        print(f'  Row {row_idx}: ' + ' | '.join(cells))
